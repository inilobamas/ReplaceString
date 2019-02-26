from flask import Flask, render_template, request
from docx import Document
import os, shutil
 
app = Flask(__name__)
 
@app.route("/")
def index():
    return render_template('input_data.html')
 
@app.route("/convert", methods=['POST'])
def convert():
    name = request.form['name']
    address = request.form['address']

    #TODO ganti path
    ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
    filePath = os.path.join(ROOT_DIR, 'source')

    for dname, dirs, files in os.walk(filePath):
        for fname in files:
            print("fname:",fname)
            folderPath = os.path.join(ROOT_DIR, 'export')
            filename = filePath + "/" + fname
            shutil.copy(filename, folderPath)
            doc_path = os.path.join(folderPath, fname)

            document = Document(filename)

            replacements = {
                '{{nama}}': name,
                '{{alamat}}': address,
                '{{pemasukan1}}': '1000',
                '{{pengeluaran1}}': '500'
            }

            for p in document.paragraphs:
                for key in replacements:
                    if key in p.text:
                        inline = p.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                text = inline[i].text.replace(key, replacements[key])
                                inline[i].text = text

            for table in document.tables:
                print('table:', table)
                for cell in table.cells:
                    print('cell:', cell)
                    for paragraph in cell.paragraphs:
                        print("paragraph:", paragraph.text)
                        for key_table in replacements:
                            print("key_table:", key_table)
                            print("replacements[key_table]:", replacements[key_table])
                            # if key_table in paragraph.text:
                            #     inline_table = paragraph.runs
                            #     # Loop added to work with runs (strings with same style)
                            #     for i in range(len(inline_table)):
                            #         if key_table in inline_table[i].text:
                            #             text = inline_table[i].text.replace(key_table, replacements[key_table])
                            #             inline_table[i].text = text

            # for table in document.tables:
            #     for cell in table.cells:
            #         for paragraph in cell.paragraphs:
            #             for key in replacements:
            #                 if key in paragraph.text:
            #                     inline = paragraph.runs
            #                     # Loop added to work with runs (strings with same style)
            #                     for i in range(len(inline)):
            #                         text = inline[i].text.replace(key, replacements[key])
            #                         inline[i].text = text

            # for table in document.tables:
            #     for row in table.rows:
            #         for cell in row.cells:
            #             inline = cell.runs
            #             if '{{pemasukan1}}' in cell.text:
            #                 for i in range(len(inline)):
            #                     if '{{pemasukan1}}' in inline[i].text:
            #                         text = inline[i].text.replace('{{pemasukan1}}', 'pemasukanku')
            #                         inline[i].text = text
            #             if '{{pengeluaran1}}' in cell.text:
            #                 for i in range(len(inline)):
            #                     if '{{pengeluaran1}}' in inline[i].text:
            #                         text = inline[i].text.replace('{{pengeluaran1}}', 'pengeluaranku')
            #                         inline[i].text = text


                        # inline = cell.runs
                        # # Loop added to work with runs (strings with same style)
                        # for i in range(len(inline)):
                        #     if '{{pemasukan1}}' in cell.text:
                        #         if cell.text in inline[i].text:
                        #             text = inline[i].text.replace(cell.text, "pemasukanku")
                        #             inline[i].text = text
                        #     if '{{pengeluaran1}}' in cell.text:
                        #         if cell.text in inline[i].text:
                        #             text = inline[i].text.replace(cell.text, "pengeluaranku")
                        #             inline[i].text = text


                        # print(cell.text)
                        # if '{{pemasukan1}}' in cell.text:
                        #     cell.text = '1000'
                        # if '{{pengeluaran1}}' in cell.text:
                        #     cell.text = '100'

            document.save(doc_path)

    return 'Hello %s your address is %s <br/> <a href="/">Back Home</a>' % (name, address)
 
if __name__ == "__main__":
    app.run()